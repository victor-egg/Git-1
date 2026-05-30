import os
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn

class DocFormatter:
    def __init__(self):
        self.doc = None
        self.file_path = None
        self.is_modified = False

    def load_document(self):
        path = input("请输入文档路径（支持.docx）: ").strip()
        if not os.path.exists(path):
            print("❌ 错误：文件路径不存在，请检查后重试。")
            return
        if not path.lower().endswith('.docx'):
            print("❌ 错误：暂不支持该文件格式，请使用.docx文件。")
            return
        try:
            self.doc = Document(path)
            self.file_path = path
            self.is_modified = False
            print(f"✅ 成功读取文档：{os.path.basename(path)}")
        except Exception as e:
            print(f"❌ 错误：读取文档失败 - {e}")

    def clean_and_format(self):
        if self.doc is None:
            print("❌ 错误：请先读取文档！")
            return
        print("⏳ 正在执行冗余清理与格式统一设置...")
        for para in self.doc.paragraphs:
            original_text = para.text
            # 1. 冗余格式清理：清理连续空格、首尾空格
            cleaned_text = re.sub(r'\s+', ' ', original_text).strip()
            if not cleaned_text:
                continue  # 跳过空段落（相当于清理连续空行）

            # 直接替换段落文本（因后续会统一重设格式，此操作安全高效）
            para.text = cleaned_text

            # 2. 统一格式设置：固定行距18磅
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            para.paragraph_format.line_spacing = Pt(18)

            # 设置字体：宋体，小四（12磅）
            for run in para.runs:
                run.font.name = '宋体'
                run.font.size = Pt(12)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        self.is_modified = True
        print("✅ 格式排版与清理完成！")

    def count_words(self):
        """进阶功能1：文档字数统计"""
        if self.doc is None:
            print("❌ 错误：请先读取文档！")
            return
        
        total_chars = 0
        for para in self.doc.paragraphs:
            # 统计去除常见空白符（空格、换行、制表符、全角空格）后的有效字符数
            text = para.text.replace(" ", "").replace("\n", "").replace("\t", "").replace("\u3000", "")
            total_chars += len(text)
            
        print(f"📊 统计结果：文档有效字符总数约为 {total_chars} 字。")

    def custom_format(self):
        """进阶功能2：自定义正文格式参数"""
        if self.doc is None:
            print("❌ 错误：请先读取文档！")
            return
            
        print("\n--- 自定义正文格式设置 ---")
        font_name = input("请输入正文字体名称（如：黑体、楷体、宋体，直接回车默认宋体）: ").strip()
        if not font_name:
            font_name = '宋体'
            
        # 异常捕获：防止用户输入非数字导致程序崩溃
        while True:
            font_size_str = input("请输入正文字号大小（数字，单位：磅，如：12、14、16，直接回车默认12）: ").strip()
            if not font_size_str:
                font_size = 12.0
                break
            try:
                font_size = float(font_size_str)
                break
            except ValueError:
                print("❌ 输入无效，请输入数字！")
                
        while True:
            line_spacing_str = input("请输入固定行距（数字，单位：磅，如：18、20、22，直接回车默认18）: ").strip()
            if not line_spacing_str:
                line_spacing = 18.0
                break
            try:
                line_spacing = float(line_spacing_str)
                break
            except ValueError:
                print("❌ 输入无效，请输入数字！")
                
        print(f"⏳ 正在应用自定义格式：{font_name}，{font_size}磅，行距{line_spacing}磅...")
        
        for para in self.doc.paragraphs:
            if not para.text.strip():
                continue  # 跳过空段落
                
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            para.paragraph_format.line_spacing = Pt(line_spacing)
            
            for run in para.runs:
                run.font.name = font_name
                run.font.size = Pt(font_size)
                # 强制绑定中文字体属性，确保在Word中正确生效
                run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                
        self.is_modified = True
        print("✅ 自定义格式应用完成！")

    def save_document(self):
        if self.doc is None:
            print("❌ 错误：没有可保存的文档。")
            return
        save_path = input("请输入保存路径（默认覆盖原文件，直接回车确认）: ").strip()
        if not save_path:
            save_path = self.file_path
        if not save_path.lower().endswith('.docx'):
            save_path += '.docx'
        try:
            self.doc.save(save_path)
            self.is_modified = False
            print(f"✅ 文档已成功保存至：{save_path}")
        except Exception as e:
            print(f"❌ 错误：保存失败 - {e}")

    def run(self):
        print("=================================")
        print("欢迎使用文档自动排版工具")
        print("=================================")
        while True:
            print("\n请选择操作：")
            print("1. 读取文档文件")
            print("2. 基础格式排版（默认格式、冗余清理）")
            print("3. 统计文档字数")
            print("4. 自定义正文格式")
            print("5. 保存排版后文档")
            print("6. 退出系统")
            choice = input("请输入选项序号: ").strip()

            if choice == '1':
                self.load_document()
            elif choice == '2':
                self.clean_and_format()
            elif choice == '3':
                self.count_words()
            elif choice == '4':
                self.custom_format()
            elif choice == '5':
                self.save_document()
            elif choice == '6':
                if self.is_modified:
                    confirm = input("⚠️ 文档尚未保存，是否保存后退出？(y/n): ").strip().lower()
                    if confirm == 'y':
                        self.save_document()
                print("👋 感谢使用，再见！")
                break
            else:
                print("❌ 错误：无效的菜单选项，请输入1-6之间的数字。")

if __name__ == "__main__":
    app = DocFormatter()
    app.run()